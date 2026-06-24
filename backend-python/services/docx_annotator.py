"""
docx_annotator.py - Anotador de resultados de auditoría en el documento Word.

Versión 3.0 (Híbrido inteligente):
- Mantiene la funcionalidad Turnitin-style (comentarios sobrescritos en el margen)
- Límite aumentado a 1500 comentarios (antes 400)
- Consolida repeticiones masivas (>5 mismos errores consecutivos → 1 comentario resumen)
- Modo híbrido por severidad:
    * error      → globo de comentario + resaltado ROJO (alta visibilidad)
    * warning    → solo resaltado AMARILLO lateral (sin globo, sin ruido)
    * observation→ solo resaltado TURQUESA tenue (sin globo)
- Respeta párrafos con fórmulas OMML (no resalta texto matemático como error textual)
"""
from docx import Document
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import os
import re
import zipfile
import traceback
import io

# ── Configuración global del anotador ─────────────────────────────────────
MAX_COMMENTS = 1500            # Antes: 400. Tesis grandes ahora cubiertas completas
CONSOLIDATE_THRESHOLD = 5      # Mismo (rule, status) consecutivos → consolidar
MAX_REPORT_DETAIL = 120        # Líneas detalladas en el reporte ejecutivo
# Categorías que prioritariamente reciben globo aunque sean warning
PRIORITY_CATS = (
    "ANEXOS OBLIGATORIOS", "ESTRUCTURA", "JERARQUÍA", "JERARQUIA",
    "TABLAS", "FIGURAS", "PORTADA", "REFERENCIAS"
)


class DocxAnnotator:
    """Motor de Anotación RepoStyle 3.0 - Híbrido inteligente."""

    def __init__(self, original_path):
        self.original_path = original_path
        self.doc = Document(original_path)
        self._comment_id = 100
        self._all_paragraphs = self._map_all_paragraphs()

    def _map_all_paragraphs(self):
        """Extrae TODOS los párrafos (incluyendo tablas) en orden XML."""
        p_elements = self.doc.element.xpath('//w:p')
        paragraphs = []
        for p in p_elements:
            p_obj = Paragraph(p, self.doc._parent)
            p_obj._cached_text = p_obj.text
            paragraphs.append(p_obj)
        return paragraphs

    def annotate(self, audit_results):
        try:
            results = audit_results.get("results", [])
            stats = audit_results.get("stats", {})

            self._ensure_comments_part()

            # 1. Extraer regla global de numeración de líneas (no asociada a párrafo)
            # Se procesa aparte (solo globo, sin resaltado) y se excluye del pipeline normal
            line_numbering_rule = None
            remaining = []
            for r in results:
                rule = (r.get("rule") or "").upper()
                rule_norm = rule.replace('Á','A').replace('É','E').replace('Í','I').replace('Ó','O').replace('Ú','U')
                is_line_rule = "NUMERACION" in rule_norm and "LINEA" in rule_norm
                if is_line_rule and r.get("status") != "passed":
                    line_numbering_rule = r
                    continue
                remaining.append(r)
            results = remaining

            # 2. Filtrar solo no-pasados
            failed = [r for r in results if r.get("status") != "passed"]

            # 3. Consolidar repeticiones masivas (mismas reglas consecutivas)
            consolidated = self._consolidate_repeats(failed)

            # 4. AGRUPAR POR PÁRRAFO — todas las observaciones del mismo párrafo
            # se combinan en UN SOLO comentario para reducir ruido.
            # Esto evita que un mismo párrafo reciba 2-3 comentarios separados
            # (ej: "sin negrita" + "sin tabulación" + "sangría incorrecta").
            grouped_by_paragraph = {}     # id(target_p) → lista de resultados
            target_paragraphs = {}         # id(target_p) → objeto párrafo
            ungrouped = []                 # resultados sin párrafo asociable

            for res in consolidated:
                target = self._find_paragraph(res)
                if target is None:
                    ungrouped.append(res)
                    continue
                if self._is_math_paragraph(target):
                    continue  # No anotar fórmulas como texto
                p_id = id(target)
                grouped_by_paragraph.setdefault(p_id, []).append(res)
                target_paragraphs[p_id] = target

            # 5. Ordenar grupos por prioridad: error > warning > observation
            # El grupo entero hereda la severidad MÁS GRAVE de sus miembros
            def group_priority(items):
                statuses = [r.get("status") for r in items]
                rank = {"error": 0, "warning": 1, "observation": 2}
                worst = min((rank.get(s, 3) for s in statuses), default=3)
                page = min((r.get("page") or 9999 for r in items), default=9999)
                # Categorías críticas reciben prioridad extra
                cats = " ".join((r.get("category") or "").upper() for r in items)
                is_critical = any(c in cats for c in PRIORITY_CATS)
                return (0 if is_critical else 1, worst, page)

            sorted_groups = sorted(
                grouped_by_paragraph.items(),
                key=lambda kv: group_priority(kv[1]),
            )

            # 6. Procesar cada GRUPO de observaciones por párrafo
            annotated_count = 0
            tables_marked = set()  # Evita marcar la misma tabla varias veces
            for p_id, observations in sorted_groups:
                if annotated_count >= MAX_COMMENTS:
                    break
                target = target_paragraphs[p_id]

                # Deduplicar observaciones idénticas (misma regla + status)
                seen_rules = set()
                unique_obs = []
                for o in observations:
                    key = (o.get("rule", ""), o.get("status", ""))
                    if key in seen_rules:
                        continue
                    seen_rules.add(key)
                    unique_obs.append(o)

                # Determinar severidad más grave del grupo
                worst_status = "observation"
                for o in unique_obs:
                    s = o.get("status", "observation")
                    if s == "error":
                        worst_status = "error"
                        break
                    if s == "warning" and worst_status != "error":
                        worst_status = "warning"

                # Determinar si todas son ortografía (caso especial: verde)
                cats = [(o.get("category") or "").upper() for o in unique_obs]
                all_ortografia = all("ORTOGRAF" in c for c in cats)
                any_priority = any(
                    any(p in c for p in PRIORITY_CATS) for c in cats
                )

                try:
                    # ── Si es un grupo de errores de TABLA, marcar la tabla completa ──
                    is_table_group = any("TABLAS" in c for c in cats) and self._is_in_table(target)
                    if is_table_group:
                        tbl_el = self._get_table_ancestor(target)
                        if tbl_el is not None:
                            tbl_key = id(tbl_el)
                            if tbl_key not in tables_marked:
                                tables_marked.add(tbl_key)
                                self._apply_table_level_marker(tbl_el, unique_obs, worst_status)
                                annotated_count += 1
                            continue  # Saltar resaltado de celda individual

                    # ── Procesamiento normal (no tabla) ──
                    # Elegir color según severidad más grave
                    if all_ortografia:
                        color = None
                        needs_comment = False
                    elif worst_status == "error":
                        color = WD_COLOR_INDEX.RED
                        needs_comment = True
                    elif worst_status == "warning":
                        color = WD_COLOR_INDEX.YELLOW
                        needs_comment = any_priority
                    else:
                        # observation / info / passed → sin resaltado
                        color = None
                        needs_comment = False

                    if color is not None:
                        self._apply_highlight(target, color)

                    if needs_comment:
                        comment_text = self._build_grouped_comment_text(unique_obs)
                        self._add_native_comment(target, comment_text)
                        annotated_count += 1
                except Exception as e:
                    print(f"Error anotando grupo: {e}")

            # 7. Marcador global de numeración de líneas (solo globo, sin resaltado)
            if line_numbering_rule is not None:
                self._apply_global_line_numbering_marker(line_numbering_rule)

            # 8. Reporte ejecutivo al inicio
            self._add_institutional_report(stats, results)

            out_name = f"auditado_{os.path.basename(self.original_path)}"
            out_path = os.path.join(os.path.dirname(self.original_path), out_name)
            self.doc.save(out_path)

            # 9. Resaltar números de página en footers (rojo si hay error de paginación)
            try:
                has_page_errors = any(
                    r.get("status") == "error" and "pag" in (r.get("category") or "").lower()
                    for r in results
                )
                if has_page_errors:
                    self._highlight_footer_page_numbers(out_path)
            except Exception as footer_err:
                print(f"⚠️ Error resaltando footers: {footer_err}")
            return out_path, out_name
        except Exception as e:
            print(f"🛑 Error en Anotador: {str(e)}")
            traceback.print_exc()
            return self.original_path, os.path.basename(self.original_path)

    # ── Consolidación de repeticiones ─────────────────────────────────────

    def _consolidate_repeats(self, failed):
        """
        Cuando una misma regla con el mismo status aparece N veces consecutivas
        (N >= CONSOLIDATE_THRESHOLD), agrupa la primera con un comentario
        resumen y degrada el resto a mismo status pero con marca de "agrupado"
        para que reciban solo resaltado, no globo.

        Esto reduce el ruido (ej: 50 comentarios idénticos de "Interlineado 1.5"
        → 1 comentario "Se detectaron 50 párrafos consecutivos con...").
        """
        if not failed:
            return failed

        out = []
        i = 0
        n = len(failed)
        while i < n:
            r = failed[i]
            rule = r.get("rule", "")
            status = r.get("status", "")
            # Buscar grupo consecutivo
            j = i + 1
            while j < n and failed[j].get("rule") == rule and failed[j].get("status") == status:
                j += 1
            group_size = j - i

            if group_size >= CONSOLIDATE_THRESHOLD:
                # Tomar el primero como representante del grupo, con mensaje agrupado
                first = dict(r)
                first["message"] = (
                    f"Se detectaron {group_size} párrafos consecutivos con esta misma "
                    f"observación.\n\n--- Detalle del primero ---\n{r.get('message', '')}"
                )
                first["rule"] = f"{rule} (x{group_size} consecutivos)"
                out.append(first)
                # El resto: solo resaltado, sin globo (marcamos _grouped)
                for k in range(i + 1, j):
                    g = dict(failed[k])
                    g["_grouped"] = True
                    # Forzar a observation para que solo reciba color, no globo
                    g["status"] = "observation"
                    out.append(g)
            else:
                # Sin grupo grande: agregar tal cual
                for k in range(i, j):
                    out.append(failed[k])
            i = j
        return out

    # ── Prioridad y filtros ───────────────────────────────────────────────

    def _priority_key(self, res):
        """Errores antes que warnings; categorías críticas primero."""
        status = res.get("status", "")
        cat = (res.get("category") or "").upper()
        status_rank = {"error": 0, "warning": 1}.get(status, 2)
        is_critical = any(c in cat for c in PRIORITY_CATS)
        return (0 if is_critical else 1, status_rank, res.get("page") or 9999)

    def _apply_global_line_numbering_marker(self, rule_data):
        """
        Inserta UN SOLO comentario globo en el párrafo donde inicia la sección
        con numeración de líneas activa, y añade una BARRA VERTICAL ROJA al
        costado izquierdo del párrafo para señalar que el problema está en el
        MARGEN (números de línea), no en el contenido del texto.
        """
        try:
            p_idx = rule_data.get("paragraphIndex")
            target = None
            if p_idx is not None and p_idx < len(self._all_paragraphs):
                target = self._all_paragraphs[p_idx]

            if target is None:
                for p in self._all_paragraphs:
                    p_text = (getattr(p, '_cached_text', None) or p.text or '').strip()
                    if p_text:
                        target = p
                        break

            if target is not None:
                # ── Barra vertical roja al costado IZQUIERDO del párrafo ──
                # Señala visualmente "el problema está en el margen izquierdo"
                # (donde Word pinta los números de línea). No modifica el texto.
                pPr = target._element.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    target._element.insert(0, pPr)
                pBdr = pPr.find(qn('w:pBdr'))
                if pBdr is None:
                    pBdr = OxmlElement('w:pBdr')
                    pPr.append(pBdr)
                left_border = OxmlElement('w:left')
                left_border.set(qn('w:val'), 'single')
                left_border.set(qn('w:sz'), '18')
                left_border.set(qn('w:space'), '8')
                left_border.set(qn('w:color'), 'CC0000')
                pBdr.append(left_border)

                # ── Comentario explícito sobre los números de línea ──
                comment_text = (
                    "NUMERACIÓN SUCESIVA DE LÍNEA ACTIVA\n"
                    "\n"
                    "Este documento tiene números de línea en el margen "
                    "(1, 2, 3...). Esos números solo se permiten en "
                    "BORRADORES de revisión, NUNCA en la versión final.\n"
                    "\n"
                    "SOLUCIÓN: Diseño → Números de línea → Ninguno."
                )
                self._add_native_comment(target, comment_text)
        except Exception as e:
            print(f"Error aplicando marcador de numeración de líneas: {e}")

    def _is_in_table(self, paragraph):
        """Detecta si un párrafo pertenece a una tabla."""
        try:
            for ancestor in paragraph._element.iterancestors():
                tag = ancestor.tag
                if isinstance(tag, str) and tag.endswith('}tbl'):
                    return True
        except Exception:
            pass
        return False

    def _get_table_ancestor(self, paragraph):
        """Retorna el elemento w:tbl ancestro del párrafo, o None."""
        try:
            for ancestor in paragraph._element.iterancestors():
                tag = ancestor.tag
                if isinstance(tag, str) and tag.endswith('}tbl'):
                    return ancestor
        except Exception:
            pass
        return None

    def _get_first_paragraph_in_table(self, tbl_element):
        """Retorna el primer Paragraph con texto dentro de la tabla."""
        try:
            for p_el in tbl_element.iter():
                tag = p_el.tag
                if isinstance(tag, str) and tag.endswith('}p'):
                    from docx.text.paragraph import Paragraph
                    p_obj = Paragraph(p_el, self.doc._parent)
                    if (getattr(p_obj, '_cached_text', None) or p_obj.text or '').strip():
                        return p_obj
        except Exception:
            pass
        return None

    def _apply_table_level_marker(self, tbl_element, observations, worst_status):
        """
        Aplica un marcador visual a TODA la tabla (borde exterior coloreado)
        y añade un comentario en el primer párrafo de la tabla.
        No resalta celdas individuales.
        """
        try:
            # Borde exterior de la tabla según severidad
            border_color = 'FF0000' if worst_status == 'error' else ('FFC000' if worst_status == 'warning' else '00BFFF')
            tbl_pr = tbl_element.find(qn('w:tblPr'))
            if tbl_pr is None:
                tbl_pr = OxmlElement('w:tblPr')
                tbl_element.insert(0, tbl_pr)
            tbl_borders = OxmlElement('w:tblBorders')
            for side in ('top', 'left', 'bottom', 'right'):
                border_el = OxmlElement(f'w:{side}')
                border_el.set(qn('w:val'), 'single')
                border_el.set(qn('w:sz'), '12')
                border_el.set(qn('w:space'), '4')
                border_el.set(qn('w:color'), border_color)
                tbl_borders.append(border_el)
            tbl_pr.append(tbl_borders)

            # Comentario en el primer párrafo de la tabla
            first_p = self._get_first_paragraph_in_table(tbl_element)
            if first_p is not None:
                comment_text = self._build_grouped_comment_text(observations)
                self._add_native_comment(first_p, comment_text)
        except Exception as e:
            print(f"Error en marcador de tabla: {e}")

    def _is_math_paragraph(self, paragraph):
        """Detecta si el párrafo contiene fórmulas OMML (oMath/oMathPara)."""
        try:
            math_nodes = paragraph._element.xpath(
                './/*[local-name()="oMath" or local-name()="oMathPara"]'
            )
            return len(math_nodes) > 0
        except Exception:
            return False

    # ── Construcción del texto del comentario ─────────────────────────────

    def _build_comment_text(self, res):
        """
        Genera el texto del comentario en formato formal, sin iconos.
        Estructura clara y legible para revisión académica.
        """
        status = (res.get("status") or "").upper()
        severity_label = {
            "ERROR": "Severidad: Error crítico",
            "WARNING": "Severidad: Advertencia",
            "OBSERVATION": "Severidad: Observación",
        }.get(status, "Severidad: Observación")
        section = res.get("section") or "General"

        return (
            f"OBSERVACIÓN DE FORMATO\n"
            f"\n"
            f"**Regla:** {res.get('rule')}\n"
            f"**Categoría:** {res.get('category')}\n"
            f"**Sección:** {section}\n"
            f"**{severity_label}**\n"
            f"\n"
            f"**Hallado:** {res.get('actual')}\n"
            f"**Requerido:** {res.get('expected')}\n"
            f"\n"
            f"**Descripción:**\n"
            f"{res.get('message')}"
        )

    def _build_grouped_comment_text(self, observations):
        """
        Genera UN SOLO comentario que combina TODAS las observaciones del
        mismo párrafo. Si hay una sola observación, retorna el formato
        estándar. Si hay varias, lista cada una numerada.
        """
        if len(observations) == 1:
            return self._build_comment_text(observations[0])

        # Determinar sección (tomar del primer resultado con datos)
        section = None
        for o in observations:
            if section is None and o.get("section"):
                section = o.get("section")
        section = section or "General"

        # Determinar severidad global
        rank = {"error": 0, "warning": 1, "observation": 2}
        worst = min(
            (rank.get(o.get("status"), 3) for o in observations),
            default=3,
        )
        global_severity = {
            0: "Error crítico",
            1: "Advertencia",
            2: "Observación",
        }.get(worst, "Observación")

        # Encabezado
        n = len(observations)
        header = (
            f"OBSERVACIÓN DE FORMATO MÚLTIPLE\n"
            f"\n"
            f"Este párrafo tiene {n} observaciones de formato.\n"
            f"**Severidad más grave del grupo:** {global_severity}\n"
            f"**Sección:** {section}\n"
            f"\n"
            f"────────────────────────────────────────\n"
        )

        # Listar cada observación
        bodies = []
        for i, o in enumerate(observations, 1):
            sev = (o.get("status") or "").upper()
            sev_label = {
                "ERROR": "Error crítico",
                "WARNING": "Advertencia",
                "OBSERVATION": "Observación",
            }.get(sev, "Observación")
            bodies.append(
                f"\n**[{i}] {o.get('rule')}**\n"
                f"    Severidad: {sev_label}\n"
                f"    Categoría: {o.get('category')}\n"
                f"    Hallado: {o.get('actual')}\n"
                f"    Requerido: {o.get('expected')}\n"
                f"    Descripción: {o.get('message')}\n"
            )

        return header + "".join(bodies) + (
            f"\n────────────────────────────────────────\n"
            f"Total: {n} observaciones en este párrafo."
        )

    # ── Aplicación de resaltado al párrafo ────────────────────────────────

    def _apply_highlight(self, target, color):
        """Aplica color de resaltado a todos los runs del párrafo (incluyendo hyperlinks y dibujos)."""
        r_elements = target._element.xpath('.//w:r')
        if r_elements:
            from docx.text.run import Run
            for r_el in r_elements:
                run_obj = Run(r_el, target)
                run_obj.font.highlight_color = color
        else:
            new_run = target.add_run(" ")
            new_run.font.highlight_color = color

    # ── Localización de párrafos ──────────────────────────────────────────

    def _find_paragraph(self, res):
        p_text = res.get("paragraphText")
        p_idx = res.get("paragraphIndex")

        # 1. Sincronización por Índice XML (Máxima Precisión)
        if p_idx is not None and p_idx < len(self._all_paragraphs):
            p = self._all_paragraphs[p_idx]
            p_txt_val = getattr(p, '_cached_text', None) or p.text
            if not p_text or str(p_text)[:15] in p_txt_val or str(p_text).startswith("[Imagen"):
                return p

        # 2. Fallback: Búsqueda por Texto cerca de p_idx
        if p_text and len(str(p_text).strip()) > 5:
            search_txt = str(p_text).strip()
            if p_idx is not None:
                start_idx = max(0, p_idx - 50)
                end_idx = min(len(self._all_paragraphs), p_idx + 50)
                search_range = self._all_paragraphs[start_idx:end_idx]
            else:
                search_range = self._all_paragraphs

            for p in search_range:
                p_txt_val = getattr(p, '_cached_text', None) or p.text
                if search_txt in p_txt_val:
                    return p
        return None

    # ── Infraestructura OOXML de comentarios ──────────────────────────────

    def _ensure_comments_part(self):
        try:
            _ = self.doc.part.comments
        except Exception:
            pass

    def _add_native_comment(self, paragraph, text, author='RepoStyle', initials='RS'):
        """Añade un comentario nativo de Word que aparece en el margen derecho."""
        try:
            comments_part = self.doc.part._comments_part
            comments_xml = comments_part.element

            cid = str(self._comment_id)
            self._comment_id += 1

            comment = OxmlElement('w:comment')
            comment.set(qn('w:id'), cid)
            comment.set(qn('w:author'), author)
            comment.set(qn('w:initials'), initials)

            lines = text.split('\n')
            for line in lines:
                p_line = OxmlElement('w:p')
                parts = line.split('**')
                is_bold = False
                for part in parts:
                    r = OxmlElement('w:r')
                    if is_bold:
                        rPr = OxmlElement('w:rPr')
                        b = OxmlElement('w:b')
                        rPr.append(b)
                        r.append(rPr)
                    t = OxmlElement('w:t')
                    t.set(qn('xml:space'), 'preserve')
                    t.text = part
                    r.append(t)
                    p_line.append(r)
                    is_bold = not is_bold
                comment.append(p_line)
            comments_xml.append(comment)

            # Marcar el rango en el documento
            pPr = paragraph._element.find(qn('w:pPr'))
            insert_idx = paragraph._element.index(pPr) + 1 if pPr is not None else 0

            range_start = OxmlElement('w:commentRangeStart')
            range_start.set(qn('w:id'), cid)
            paragraph._element.insert(insert_idx, range_start)

            range_end = OxmlElement('w:commentRangeEnd')
            range_end.set(qn('w:id'), cid)
            paragraph._element.append(range_end)

            # Referencia visual (globo)
            ref_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rStyle = OxmlElement('w:rStyle')
            rStyle.set(qn('w:val'), 'CommentReference')
            rPr.append(rStyle)
            ref_run.append(rPr)

            comment_ref = OxmlElement('w:commentReference')
            comment_ref.set(qn('w:id'), cid)
            ref_run.append(comment_ref)
            paragraph._element.append(ref_run)

        except Exception as e:
            print(f"Error en comentario nativo: {e}")

    # ── Reporte ejecutivo institucional ───────────────────────────────────

    def _add_institutional_report(self, stats, results):
        """Reporte ejecutivo institucional, formato formal sin iconos."""
        if not self.doc.paragraphs:
            self.doc.add_paragraph("")
        first = self.doc.paragraphs[0]

        t1 = first.insert_paragraph_before("REPORTE INSTITUCIONAL DE AUDITORÍA DE TESIS")
        t1.runs[0].bold = True
        t1.runs[0].font.size = Pt(18)
        t1.runs[0].font.color.rgb = RGBColor(0, 51, 102)

        t2 = first.insert_paragraph_before(
            "Vicerrectorado de Investigación — Universidad Nacional del Altiplano de Puno"
        )
        t2.runs[0].bold = True
        t2.runs[0].font.size = Pt(11)
        t2.runs[0].font.color.rgb = RGBColor(80, 80, 80)

        first.insert_paragraph_before("")  # separador
        p_score = first.insert_paragraph_before(
            f"Puntaje de cumplimiento: {stats.get('score', 0)} / 100"
        )
        p_score.runs[0].bold = True
        p_score.runs[0].font.size = Pt(14)
        p_score.runs[0].font.color.rgb = RGBColor(0, 0, 0)

        stats_p = first.insert_paragraph_before(
            f"Errores críticos: {stats.get('errors', 0)}      "
            f"Advertencias: {stats.get('warnings', 0)}      "
            f"Validaciones aprobadas: {stats.get('passed', 0)}"
        )
        stats_p.runs[0].font.size = Pt(11)

        # Leyenda del modo híbrido
        leg = first.insert_paragraph_before(
            "Leyenda de marcado: Resaltado rojo indica error crítico (con comentario al margen). "
            "Resaltado amarillo indica advertencia individual. Resaltado verde claro indica "
            "sugerencia ortográfica. Resaltado turquesa indica observación menor. "
            "Resaltado azul indica marcador GLOBAL del documento (ej: numeración de líneas activa) "
            "— en este caso hay UN solo comentario al inicio explicando la causa, no uno por párrafo."
        )
        leg.runs[0].font.size = Pt(9)
        leg.runs[0].italic = True
        leg.runs[0].font.color.rgb = RGBColor(100, 100, 100)

        first.insert_paragraph_before("")  # separador
        header = first.insert_paragraph_before("DETALLE DE OBSERVACIONES")
        header.runs[0].bold = True
        header.runs[0].font.size = Pt(13)
        header.runs[0].font.color.rgb = RGBColor(0, 51, 102)

        sub = first.insert_paragraph_before("Ordenadas por página de aparición y sección.")
        sub.runs[0].font.size = Pt(10)
        sub.runs[0].italic = True
        sub.runs[0].font.color.rgb = RGBColor(100, 100, 100)

        failed = [r for r in results if r.get("status") != "passed"]
        failed_sorted = sorted(failed, key=lambda r: (r.get("page") or 9999, r.get("section") or ""))
        failed_display = failed_sorted[:MAX_REPORT_DETAIL]

        idx = 1
        last_section = None
        for res in failed_display:
            section = res.get("section") or "General"
            page = res.get("page")
            page_label = f"Página {page}" if page else "Página no determinada"

            if section != last_section:
                sec_p = first.insert_paragraph_before(f"Sección: {section}")
                sec_p.runs[0].bold = True
                sec_p.runs[0].font.size = Pt(11)
                sec_p.runs[0].font.color.rgb = RGBColor(0, 70, 127)
                last_section = section

            severity = "Error" if res.get("status") == "error" else "Advertencia"
            p = first.insert_paragraph_before(
                f"  {idx}. [{severity}] [{page_label}] [{res.get('category')}] {res.get('rule')}"
            )
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.color.rgb = (
                RGBColor(180, 0, 0) if res.get("status") == "error" else RGBColor(150, 100, 0)
            )

            det_p = first.insert_paragraph_before(
                f"      Hallado: {res.get('actual')}   |   Esperado: {res.get('expected')}"
            )
            det_p.runs[0].font.size = Pt(10)
            det_p.runs[0].font.color.rgb = RGBColor(60, 60, 60)

            p_text = res.get("paragraphText") or ""
            if p_text and p_text.strip():
                snip = p_text.strip()[:80] + ("..." if len(p_text.strip()) > 80 else "")
                ctx_p = first.insert_paragraph_before(f"      Texto referencial: \"{snip}\"")
                ctx_p.runs[0].font.size = Pt(9)
                ctx_p.runs[0].font.color.rgb = RGBColor(120, 120, 120)
                ctx_p.runs[0].italic = True

            idx += 1

        if len(failed_sorted) > MAX_REPORT_DETAIL:
            extra_count = len(failed_sorted) - MAX_REPORT_DETAIL
            p_extra = first.insert_paragraph_before(
                f"\nSe omitieron {extra_count} observaciones adicionales en este reporte impreso "
                f"para mantener el documento legible. Consulte la plataforma web de RepoStyle "
                f"para acceder al listado completo e interactivo."
            )
            p_extra.runs[0].italic = True
            p_extra.runs[0].font.size = Pt(10)
            p_extra.runs[0].font.color.rgb = RGBColor(0, 51, 102)

        first.insert_paragraph_before("").add_run().add_break(WD_BREAK.PAGE)

    def _highlight_footer_page_numbers(self, docx_path):
        """Resalta los números de página en los footers del documento (rojo)."""
        if not docx_path or not os.path.exists(docx_path):
            return

        try:
            from lxml import etree
            import io

            buffer = io.BytesIO()
            with open(docx_path, 'rb') as f:
                buffer.write(f.read())

            modified = False
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            with zipfile.ZipFile(buffer, 'a') as z:
                footer_files = [n for n in z.namelist() if re.match(r'word/footer\d*\.xml', n)]
                for fpath in footer_files:
                    xml_bytes = z.read(fpath)
                    root = etree.fromstring(xml_bytes)
                    doc_modified = False

                    for p in root.iter(f'{{{ns["w"]}}}p'):
                        has_page = False
                        for instr in p.iter(f'{{{ns["w"]}}}instrText'):
                            if instr.text and 'PAGE' in instr.text.upper():
                                has_page = True
                                break
                        for fld in p.iter(f'{{{ns["w"]}}}fldSimple'):
                            instr = fld.get(f'{{{ns["w"]}}}instr', '')
                            if 'PAGE' in instr.upper():
                                has_page = True
                                break

                        if not has_page:
                            continue

                        for r_el in p.iter(f'{{{ns["w"]}}}r'):
                            rpr = r_el.find(f'{{{ns["w"]}}}rPr')
                            if rpr is None:
                                rpr = etree.SubElement(r_el, f'{{{ns["w"]}}}rPr')
                            old_hl = rpr.find(f'{{{ns["w"]}}}highlight')
                            if old_hl is not None:
                                rpr.remove(old_hl)
                            hl = etree.SubElement(rpr, f'{{{ns["w"]}}}highlight')
                            hl.set(f'{{{ns["w"]}}}val', 'red')

                        doc_modified = True

                    if doc_modified:
                        modified = True
                        z.writestr(fpath, etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True))

            if modified:
                with open(docx_path, 'wb') as f:
                    f.write(buffer.getvalue())
                print("[Annotator] Footers con PAGE resaltados en ROJO.")
        except Exception as e:
            print(f"[Annotator] Error en _highlight_footer_page_numbers: {e}")
