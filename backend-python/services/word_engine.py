"""
word_engine.py - Motor principal de auditoría de tesis.

Este archivo es el ORQUESTADOR que:
1. Extrae los datos del documento Word (.docx)
2. Delega cada aspecto de la auditoría a un módulo especializado en services/auditors/
3. Consolida y devuelve los resultados

Módulos de auditoría (en services/auditors/):
- portada.py              → Portada (primera página)
- indice_general.py       → Índice General (tabla de contenidos)
- indice_tablas_figuras.py → Índices de Tablas, Figuras y Anexos
- tablas_figuras.py       → Tablas y Figuras en el cuerpo
- titulos_nivel1.py       → Títulos de Nivel 1 (Capítulos, secciones principales)
- titulos_nivel2345.py    → Títulos de Nivel 2, 3, 4 y 5
- vinetas.py              → Viñetas (todo el documento)
- contenido_parrafos.py   → Párrafos de contenido
- configuracion_pagina.py → Configuración de página y márgenes
- resumen_abstract.py     → Resumen, Abstract y Agradecimientos
- anexos.py               → Sección de Anexos
- estilo_escritura.py     → Estilo de escritura y hábitos de formato
"""
from docx import Document
import os
import re
import unicodedata
import zipfile
from lxml import etree
from utils.style_resolver import StyleResolver, parse_rpr, parse_ppr, NSMAP
from utils.converter import DocConverter
from services.linguistic_analyzer import LinguisticAnalyzer

# Importar todos los auditores modulares
from services.auditors.portada import PortadaAuditor
from services.auditors.indice_general import IndiceGeneralAuditor
from services.auditors.indice_tablas_figuras import IndiceTablasFigurasAuditor
from services.auditors.paginacion_indices import PaginacionIndicesAuditor
from services.auditors.tablas import TablasAuditor
from services.auditors.figuras import FigurasAuditor
from services.auditors.capitulo_nivel1 import CapituloNivel1Auditor
from services.auditors.capitulo_nivel2 import CapituloNivel2Auditor
from services.auditors.capitulo_nivel345 import CapituloNivel345Auditor
from services.auditors.vinetas import VinetasAuditor
from services.auditors.configuracion_pagina import ConfiguracionPaginaAuditor
from services.auditors.agradecimientos import AgradecimientosAuditor
from services.auditors.resumen import ResumenAuditor
from services.auditors.abstract import AbstractAuditor
from services.auditors.anexos import AnexosAuditor
from services.auditors.estilo_escritura import EstiloEscrituraAuditor
from services.auditors.etiquetas_jurados import EtiquetasJuradosAuditor
from services.auditors.acronimos import AcronimosAuditor
from services.auditors.reporte_similitud import ReporteSimilitudAuditor
from services.auditors.dedicatoria import DedicatoriaAuditor
from services.auditors.conclusiones_recomendaciones import ConclusionesRecomendacionesAuditor
from services.auditors.referencias_bibliograficas import ReferenciasBibliograficasAuditor
from services.auditors.secuencia_titulos import SecuenciaTitulosAuditor
from services.auditors.declaracion_autenticidad import DeclaracionAutenticidadAuditor
from services.auditors.autorizacion_deposito import AutorizacionDepositoAuditor
from services.auditors.cuadros_texto import CuadrosTextoAuditor
from services.auditors.capturas import CapturasAuditor
from services.auditors.ortografia import OrtografiaAuditor
from services.auditors.numeracion_paginas import NumeracionPaginasAuditor
from services.auditors.secuencia_tabla_figura import SecuenciaTablaFiguraAuditor

# Palabras aproximadas por página en formato tesis
_WORDS_PER_PAGE = 350


class WordAuditEngine:
    def __init__(self, file_path):
        self.original_path = file_path
        self.working_path = file_path
        self.stats = {"score": 100, "errors": 0, "warnings": 0, "passed": 0}
        self.results = []
        self.paragraphs = []
        self.sections_found = set()
        self.linguistic = LinguisticAnalyzer()
        # Contexto de posición
        self._current_section = "Inicio del Documento"
        self._word_count_accum = 0
        self.anexos_start_idx = -1

    def run_audit(self):
        try:
            output_dir = os.path.dirname(self.original_path)
            self.working_path = DocConverter.standardize_to_docx(self.original_path, output_dir)
            self.document = Document(self.working_path)
            self.resolver = StyleResolver(self.working_path)
            self._extract_data()

            # ── Ejecutar cada auditor modular ──────────────────────────
            PortadaAuditor(self).audit()              # 📄 Portada
            ReporteSimilitudAuditor(self).audit()      # 📄 Reporte de Similitud
            DedicatoriaAuditor(self).audit()          # 📄 Dedicatoria
            AgradecimientosAuditor(self).audit()       # 📋 Agradecimientos
            IndiceGeneralAuditor(self).audit()         # 🗂️ Índice General
            IndiceTablasFigurasAuditor(self).audit()   # 🗂️ Índice Tablas/Figuras
            PaginacionIndicesAuditor(self).audit()     # 🔢 Paginación de Índices
            ResumenAuditor(self).audit()               # 📋 Resumen
            AbstractAuditor(self).audit()              # 📋 Abstract
            TablasAuditor(self).audit()                # 📊 Tablas
            FigurasAuditor(self).audit()               # 📊 Figuras
            SecuenciaTablaFiguraAuditor(self).audit()  # 📊 Secuencia: etiqueta → título → contenido → nota
            CapituloNivel1Auditor(self).audit()        # 📌 Nivel 1: Títulos + Contenido
            CapituloNivel2Auditor(self).audit()        # 📌 Nivel 2: Títulos + Contenido
            CapituloNivel345Auditor(self).audit()      # 📌 Nivel 3-5: Títulos + Contenido
            SecuenciaTitulosAuditor(self).audit()      # 📌 Secuencia Lógica de Títulos
            VinetasAuditor(self).audit()               # 🔹 Viñetas (todos los niveles)
            ConfiguracionPaginaAuditor(self).audit()   # ⚙️ Configuración de página
            NumeracionPaginasAuditor(self).audit()     # 🔢 Numeración de páginas (preliminares sin número)
            EstiloEscrituraAuditor(self).audit()       # ✍️ Estilo de escritura
            CuadrosTextoAuditor(self).audit()          # 🚫 Cuadros de texto (evasión Turnitin)
            CapturasAuditor(self).audit()              # 🚫 Capturas de pantalla con texto
            OrtografiaAuditor(self).audit()            # 📝 Revisión ortográfica
            ConclusionesRecomendacionesAuditor(self).audit() # 💡 Conclusiones y Recomendaciones
            ReferenciasBibliograficasAuditor(self).audit() # 📚 Referencias Bibliográficas
            AnexosAuditor(self).audit()                # 📎 Anexos
            DeclaracionAutenticidadAuditor(self).audit() # 📜 Declaración Jurada (OBLIGATORIO UNAP)
            AutorizacionDepositoAuditor(self).audit()    # 📜 Autorización Depósito (OBLIGATORIO UNAP)
            EtiquetasJuradosAuditor(self).audit()      # 🏷️ Etiquetas Jurados
            AcronimosAuditor(self).audit()             # 🔠 Acrónimos

            return self._finalize()
        except Exception as e:
            self._add("Sistema", "Error Crítico", "error", str(e))
            return self._finalize()

    # ══════════════════════════════════════════════════════════════════════
    # EXTRACCIÓN DE DATOS DEL DOCUMENTO
    # ══════════════════════════════════════════════════════════════════════

    def _extract_data(self):
        with zipfile.ZipFile(self.working_path, 'r') as z:
            root = etree.fromstring(z.read('word/document.xml'))
        body = root.find('w:body', NSMAP)

        # Detección de numeración de líneas en secciones (borradores)
        # Solo se considera activa si lnNumType tiene countBy >= 1.
        # Muchos documentos Word tienen lnNumType residual de plantillas
        # con countBy="0" o sin atributos → no es numeración visible real.
        self.has_line_numbering = False
        sectPr_elements = root.findall('.//w:sectPr', NSMAP)
        for sectPr in sectPr_elements:
            lnNumType = sectPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lnNumType', NSMAP)
            if lnNumType is None:
                lnNumType = sectPr.find('w:lnNumType', NSMAP)
            if lnNumType is not None:
                # Verificar que countBy >= 1 (numeración realmente visible)
                count_by = lnNumType.get(f'{{{NSMAP["w"]}}}countBy')
                if count_by is None:
                    count_by = lnNumType.get('countBy')
                try:
                    count_by_val = int(count_by) if count_by else 0
                except (ValueError, TypeError):
                    count_by_val = 0
                if count_by_val >= 1:
                    self.has_line_numbering = True
                    break

        # ── Parsear numbering.xml para mapear numId e ilvl a sus formatos de viñeta ──
        bullet_definitions = {}
        try:
            with zipfile.ZipFile(self.working_path, 'r') as z:
                num_xml_data = z.read('word/numbering.xml')
                num_root = etree.fromstring(num_xml_data)
                
                # Mapear numId -> abstractNumId
                num_map = {}
                for num in num_root.findall('w:num', NSMAP):
                    numId = num.get(f'{{{NSMAP["w"]}}}numId')
                    abs_el = num.find('w:abstractNumId', NSMAP)
                    if abs_el is not None:
                        num_map[numId] = abs_el.get(f'{{{NSMAP["w"]}}}val')
                        
                # Parsear abstractNum
                abstract_nums = {}
                for absNum in num_root.findall('w:abstractNum', NSMAP):
                    absId = absNum.get(f'{{{NSMAP["w"]}}}abstractNumId')
                    levels = {}
                    for lvl in absNum.findall('w:lvl', NSMAP):
                        ilvl = lvl.get(f'{{{NSMAP["w"]}}}ilvl')
                        numFmt_el = lvl.find('w:numFmt', NSMAP)
                        numFmt = numFmt_el.get(f'{{{NSMAP["w"]}}}val') if numFmt_el is not None else 'None'
                        lvlText_el = lvl.find('w:lvlText', NSMAP)
                        lvlText = lvlText_el.get(f'{{{NSMAP["w"]}}}val') if lvlText_el is not None else ''
                        
                        rPr = lvl.find('w:rPr', NSMAP)
                        rFonts = rPr.find('w:rFonts', NSMAP) if rPr is not None else None
                        font_name = rFonts.get(f'{{{NSMAP["w"]}}}hint') or rFonts.get(f'{{{NSMAP["w"]}}}ascii') if rFonts is not None else 'None'
                        
                        levels[ilvl] = {"numFmt": numFmt, "lvlText": lvlText, "font": font_name}
                    abstract_nums[absId] = levels
                    
                # Combinar en bullet_definitions
                for numId, absId in num_map.items():
                    if absId in abstract_nums:
                        for ilvl, fmt_info in abstract_nums[absId].items():
                            bullet_definitions[(int(numId), int(ilvl))] = fmt_info
        except Exception:
            pass

        accumulated_words = 0
        current_page = 1
        page_words = 0
        current_section = "Inicio del Documento"
        in_index_zone = False
        SECTION_KEYWORDS = [
            "DEDICATORIA", "AGRADECIMIENTOS", "INDICE GENERAL",
            "INDICE DE TABLAS", "INDICE DE FIGURAS", "INDICE DE CUADROS", "INDICE DE ILUSTRACIONES",
            "INDICE DE ANEXOS", "INDICE DE GRAFICOS", "INDICE", "CONTENIDO", "TABLA DE MATERIAS",
            "TABLA DE CONTENIDOS", "TABLA DE CONTENIDO",
            "RESUMEN", "ABSTRACT", "INTRODUCCION", "CONCLUSIONES", "RECOMENDACIONES",
            "REFERENCIAS BIBLIOGRAFICAS", "ANEXOS", "DECLARACION JURADA", "AUTORIZACION PARA EL DEPOSITO",
            "ACRONIMOS", "ACRÓNIMOS"
        ]

        table_eq_cache = {}

        # ── ANÁLISIS PROFUNDO DE TABLAS ──
        # Por cada tabla del documento, capturamos:
        #   - Alineación (jc): "left", "center", "right"
        #   - Si la primera fila tiene <w:tblHeader/> (Word la repetirá en cada página)
        #   - Si la tabla cruza páginas (algún <w:r> dentro tiene lastRenderedPageBreak)
        #   - Número de filas
        # Esto permite que el auditor de tablas valide:
        #   1. La tabla debe estar alineada a IZQUIERDA (no centrada)
        #   2. Si cruza páginas → la primera fila debe tener tblHeader=true
        self.tables_info = {}
        # header_rows_by_tbl[tbl_id] = set de IDs de filas (w:tr) que son
        # encabezado (primera fila lógica o filas marcadas con <w:tblHeader/>)
        self.header_rows_by_tbl = {}
        # vmerge_continuation_cells[tbl_id] = set de IDs de celdas que son
        # continuación de un vMerge cuyo restart está en una fila de header
        self.vmerge_header_cells = {}

        w_tbl_tag = f"{{{NSMAP['w']}}}tbl"
        w_tr_tag = f"{{{NSMAP['w']}}}tr"
        w_tc_tag = f"{{{NSMAP['w']}}}tc"
        for tbl_el in body.iter(w_tbl_tag):
            tbl_id = id(tbl_el)
            tbl_pr = tbl_el.find('w:tblPr', NSMAP)
            jc_val = "left"
            if tbl_pr is not None:
                jc_el = tbl_pr.find('w:jc', NSMAP)
                if jc_el is not None:
                    jc_val = jc_el.get(f'{{{NSMAP["w"]}}}val') or "left"

            rows = list(tbl_el.iter(w_tr_tag))
            if not rows:
                continue

            # ── Detectar filas marcadas explícitamente como <w:tblHeader/> ──
            # Estas filas se repiten visualmente como encabezado en cada página.
            # Word permite tener VARIAS filas marcadas (encabezado multilínea).
            explicit_header_rows = []
            for r_idx, tr in enumerate(rows):
                tr_pr = tr.find('w:trPr', NSMAP)
                is_header = False
                if tr_pr is not None:
                    th_el = tr_pr.find('w:tblHeader', NSMAP)
                    if th_el is not None:
                        v = th_el.get(f'{{{NSMAP["w"]}}}val')
                        is_header = v is None or v not in ('0', 'false', 'off')
                if is_header:
                    explicit_header_rows.append(r_idx)

            # ── Determinar las filas que componen el encabezado visual ──
            # Si hay filas explícitas con tblHeader → esas son
            # Si no → la primera fila es el encabezado por defecto
            if explicit_header_rows:
                # Si las explícitas son consecutivas desde el inicio, usarlas
                header_row_indices = explicit_header_rows
            else:
                header_row_indices = [0]  # primera fila por defecto

            # ── Detectar celdas con vMerge="restart" en filas de encabezado ──
            # Las celdas que continúan ese merge (en filas posteriores) son
            # visualmente parte del encabezado, aunque estén vacías.
            header_row_ids = {id(rows[i]) for i in header_row_indices}
            self.header_rows_by_tbl[tbl_id] = header_row_ids

            # Mapear columnas (considerando gridSpan) a su estado de header
            # Para cada fila de header, encontrar qué celdas tienen vMerge="restart"
            vmerge_header_cell_ids = set()
            header_columns_with_vmerge = []  # lista de (row_idx, col_position)

            for h_idx in header_row_indices:
                row = rows[h_idx]
                col_pos = 0
                for tc in row.findall('w:tc', NSMAP):
                    tc_pr = tc.find('w:tcPr', NSMAP)
                    grid_span = 1
                    if tc_pr is not None:
                        gs = tc_pr.find('w:gridSpan', NSMAP)
                        if gs is not None:
                            try:
                                grid_span = int(gs.get(f'{{{NSMAP["w"]}}}val', '1'))
                            except (ValueError, TypeError):
                                grid_span = 1

                        v_merge = tc_pr.find('w:vMerge', NSMAP)
                        if v_merge is not None:
                            v = v_merge.get(f'{{{NSMAP["w"]}}}val')
                            if v == 'restart':
                                # Esta celda inicia un vMerge → continuaciones también son header
                                header_columns_with_vmerge.append((h_idx, col_pos, grid_span))
                    col_pos += grid_span

            # Buscar celdas en filas posteriores que continúen los vMerge del header
            for h_idx, col_pos, span in header_columns_with_vmerge:
                for r_idx in range(h_idx + 1, len(rows)):
                    row = rows[r_idx]
                    current_pos = 0
                    for tc in row.findall('w:tc', NSMAP):
                        tc_pr = tc.find('w:tcPr', NSMAP)
                        this_span = 1
                        is_continue = False
                        if tc_pr is not None:
                            gs = tc_pr.find('w:gridSpan', NSMAP)
                            if gs is not None:
                                try:
                                    this_span = int(gs.get(f'{{{NSMAP["w"]}}}val', '1'))
                                except (ValueError, TypeError):
                                    this_span = 1
                            v_merge = tc_pr.find('w:vMerge', NSMAP)
                            if v_merge is not None:
                                v = v_merge.get(f'{{{NSMAP["w"]}}}val')
                                is_continue = v is None or v == 'continue'

                        if current_pos == col_pos and is_continue:
                            vmerge_header_cell_ids.add(id(tc))
                        current_pos += this_span
                        if current_pos > col_pos + span:
                            break

            self.vmerge_header_cells[tbl_id] = vmerge_header_cell_ids

            # ── La tabla cruza páginas ──
            crosses_pages = bool(tbl_el.xpath('.//w:lastRenderedPageBreak', namespaces=NSMAP))

            # Texto de referencia (primera celda no vacía para identificación)
            first_cell_text = ""
            first_t = rows[0].find('.//w:t', NSMAP)
            if first_t is not None and first_t.text:
                first_cell_text = first_t.text.strip()[:30]

            self.tables_info[tbl_id] = {
                "jc": jc_val,
                "row_count": len(rows),
                "first_row_has_header": len(explicit_header_rows) > 0,
                "header_row_count": len(header_row_indices),
                "explicit_header_rows": explicit_header_rows,
                "has_merged_cells": len(vmerge_header_cell_ids) > 0 or len(header_columns_with_vmerge) > 0,
                "crosses_pages": crosses_pages,
                "first_cell_text": first_cell_text,
            }

        # ── DETECCIÓN DE CUADROS DE TEXTO (TEXT BOXES) ──
        # Los cuadros de texto en OOXML aparecen como w:txbxContent o
        # mc:AlternateContent con w:txbx. Estudiantes a veces los usan para
        # poner texto que Turnitin no detecte. La regla: NO usarlos en el cuerpo
        # de la tesis. Todo texto debe estar como párrafos normales.
        self.textboxes = []
        # Buscar txbxContent (modo moderno) y v:textbox (modo legacy)
        for txbx in body.xpath(
            './/*[local-name()="txbxContent"] | .//*[local-name()="textbox"]',
        ):
            inner_text = " ".join(
                t.text for t in txbx.iter(f'{{{NSMAP["w"]}}}t') if t.text
            ).strip()
            if not inner_text:
                continue
            # Localizar el párrafo padre más cercano (para reportar página)
            ancestor_p = next(
                (a for a in txbx.iterancestors(f'{{{NSMAP["w"]}}}p')), None
            )
            self.textboxes.append({
                "text": inner_text[:200],
                "word_count": len(inner_text.split()),
                "ancestor_p_xml_id": id(ancestor_p) if ancestor_p is not None else None,
            })

        for idx, p_el in enumerate(body.iter(f'{{{NSMAP["w"]}}}p')):
            ppr_el = p_el.find('w:pPr', NSMAP)
            explicit_ppr = parse_ppr(ppr_el)
            style_id = explicit_ppr.get('style_id', 'Normal')
            res_ppr, _ = self.resolver.resolve(style_id, explicit_ppr)

            # Detectar salto de página explícito en el XML
            _br_page = p_el.find('.//w:br[@w:type="page"]', NSMAP)
            _last_pb = p_el.find('.//w:lastRenderedPageBreak', NSMAP)
            has_page_break = (_br_page is not None) or (_last_pb is not None)

            # ── DETECCIÓN OMML (Office Math Markup Language) ──
            # Fórmulas matemáticas tienen oMath / oMathPara dentro del párrafo.
            # Marcamos el párrafo para que los auditores NO lo validen como texto normal
            # (los formatos de párrafos con fórmulas son distintos: spacing, alignment, etc.)
            has_omml = bool(
                p_el.xpath('.//*[local-name()="oMath" or local-name()="oMathPara"]')
            )
            # También detectar si el párrafo está dentro de un oMathPara (display equation)
            is_display_equation = bool(
                p_el.xpath('ancestor::*[local-name()="oMathPara"]')
            )

            # ── DETECCIÓN DE CUADROS DE TEXTO en este párrafo ──
            # Estudiantes a veces meten texto en text boxes para que Turnitin no lo
            # reconozca. Detectamos w:txbxContent (moderno) y v:textbox (legacy).
            textboxes_in_p = p_el.xpath(
                './/*[local-name()="txbxContent"] | .//*[local-name()="textbox"]'
            )
            has_textbox = len(textboxes_in_p) > 0
            textbox_text = ""
            if has_textbox:
                for tb in textboxes_in_p:
                    parts = [t.text for t in tb.iter(f'{{{NSMAP["w"]}}}t') if t.text]
                    if parts:
                        textbox_text += " ".join(parts)
                textbox_text = textbox_text.strip()[:200]

            # ── DETECCIÓN DE IMÁGENES QUE PARECEN CAPTURAS DE TEXTO ──
            # Heurística: imagen grande y horizontal (proporción > 1.5:1) en el
            # cuerpo del texto, fuera de portada/jurados/turnitin. Esto suele ser
            # una captura de pantalla con texto que el estudiante no transcribió.
            looks_like_screenshot = False
            screenshot_drawing = None
            # Esta evaluación se aplica más adelante cuando ya tenemos drawings

            # Si hay marca real de Word (lastRenderedPageBreak), usar página exacta
            has_real_page_marker = (_last_pb is not None)

            # Detectar dibujos/imagenes en el párrafo
            drawings = []
            for ext_el in p_el.xpath('.//*[local-name()="extent"]'):
                cx = ext_el.get('cx')
                cy = ext_el.get('cy')
                if cx and cy:
                    try:
                        width_cm = round(int(cx) / 360000.0, 2)
                        height_cm = round(int(cy) / 360000.0, 2)
                        drawings.append({
                            "width": width_cm,
                            "height": height_cm
                        })
                    except ValueError:
                        pass

            # Evaluar heurística de "captura de pantalla con texto":
            # Imagen grande (>10cm de ancho), proporción no cuadrada (ratio > 1.4 o < 0.7).
            # Si está cerca de una etiqueta "Figura N" → es figura legítima, no captura.
            # La evaluación final la hace el auditor (capturas.py), aquí solo marcamos
            # la propiedad para que pueda revisarse después.
            for d in drawings:
                w = d.get('width', 0)
                h = d.get('height', 0)
                if w >= 10.0 and h >= 5.0:
                    ratio = w / h if h > 0 else 1
                    if ratio > 1.5 or (ratio < 0.7 and w > 12):
                        looks_like_screenshot = True
                        screenshot_drawing = d
                        break

            txt = ""
            runs = []
            for r_el in p_el.findall('.//w:r', NSMAP):
                t = "".join([t.text for t in r_el.findall('w:t', NSMAP) if t.text])
                txt += t
                explicit_rpr = parse_rpr(r_el.find('w:rPr', NSMAP))
                _, res_rpr = self.resolver.resolve(style_id, explicit_ppr, explicit_rpr)
                runs.append({
                    "bold": res_rpr.get('bold', False), 
                    "italic": res_rpr.get('italic', False),
                    "size": res_rpr.get('font_size', 12), 
                    "font": res_rpr.get('font_name', 'Times New Roman'),
                    "text": t
                })

            word_count = len(txt.split()) if txt.strip() else 0
            accumulated_words += word_count

            # ── ESTIMACIÓN DE PÁGINA MEJORADA ──
            # Prioridad 1: lastRenderedPageBreak (marca REAL de Word del último renderizado)
            # Prioridad 2: <w:br type="page"/> (salto manual de página)
            # Prioridad 3: estimación por palabras (~350 por página)
            if has_real_page_marker:
                # Marca real → página exacta (Word renderizó aquí un salto)
                current_page += 1
                page_words = word_count
            elif has_page_break:
                current_page += 1
                page_words = word_count
            else:
                page_words += word_count
                if page_words >= _WORDS_PER_PAGE:
                    pages_to_add = page_words // _WORDS_PER_PAGE
                    current_page += pages_to_add
                    page_words = page_words % _WORDS_PER_PAGE

            estimated_page = current_page

            # Detectar si este párrafo inicia una sección conocida
            norm_txt = self._norm(txt)
            is_index_kw = norm_txt in [
                "INDICE GENERAL", "INDICE DE TABLAS", "INDICE DE FIGURAS", 
                "INDICE DE CUADROS", "INDICE DE ILUSTRACIONES", "INDICE DE ANEXOS", 
                "INDICE DE GRAFICOS", "INDICE", "CONTENIDO", "TABLA DE MATERIAS",
                "TABLA DE CONTENIDOS", "TABLA DE CONTENIDO"
            ] or (norm_txt.startswith("INDICE") and len(norm_txt) < 40)
            if is_index_kw:
                in_index_zone = True

            is_real_resumen_or_abstract = (norm_txt in ["RESUMEN", "ABSTRACT"]) and \
                                           (not "...." in txt) and \
                                           (not bool(re.search(r"\d+$", txt.strip()))) and \
                                           (res_ppr.get('alignment', 'left') == 'center')
            if is_real_resumen_or_abstract:
                in_index_zone = False

            is_index_line = "...." in txt or bool(re.search(r"\d+$", txt.strip()))
            # Salir de la zona de índice si encontramos un CAPÍTULO o INTRODUCCIÓN real del cuerpo
            if not is_index_line and ("CAPITULO" in norm_txt or "INTRODUCCION" in norm_txt):
                in_index_zone = False

            for kw in SECTION_KEYWORDS:
                if norm_txt == kw and not is_index_line:
                    current_section = txt.strip()
                    if kw == "ANEXOS" and self.anexos_start_idx == -1:
                        first_run_size = runs[0].get('size', 0) if runs else 0
                        is_centered = res_ppr.get('alignment', 'left') == 'center'
                        is_bold = any(r.get('bold') for r in runs) if runs else False
                        if 14 <= first_run_size <= 18 and is_centered and is_bold:
                            self.anexos_start_idx = idx
                    break

            if not in_index_zone:
                cap_match = re.match(r"^CAPITULO\s+(I|V|X|L|C|[0-9]+)", norm_txt)
                if cap_match:
                    current_section = txt.strip()
                    self.current_level = 1

            if not hasattr(self, 'current_level'):
                self.current_level = 1

            in_table = False
            is_table_header = False
            tbl_id_ref = None  # ID de la tabla a la que pertenece este párrafo
            w_tbl = f"{{{NSMAP['w']}}}tbl"
            w_tr = f"{{{NSMAP['w']}}}tr"
            w_tc = f"{{{NSMAP['w']}}}tc"

            tbl = next(p_el.iterancestors(w_tbl), None)
            if tbl is not None:
                tbl_id = id(tbl)
                if tbl_id in table_eq_cache:
                    is_eq = table_eq_cache[tbl_id]
                else:
                    is_eq = False
                    if tbl.xpath('.//*[local-name()="oMath" or local-name()="oMathPara"]'):
                        is_eq = True
                    else:
                        rows_iter = tbl.iter(w_tr)
                        rows = []
                        for _ in range(3):
                            try:
                                rows.append(next(rows_iter))
                            except StopIteration:
                                break
                        if len(rows) <= 2:
                            w_tc = f"{{{NSMAP['w']}}}tc"
                            w_t = f"{{{NSMAP['w']}}}t"
                            for row in rows:
                                cells = list(row.iter(w_tc))
                                if len(cells) <= 3:
                                    for cell in cells:
                                        cell_text = "".join([t.text for t in cell.iter(w_t) if t.text]).strip()
                                        if re.match(r'^\(\s*\d+(\s*\.\s*\d+)*\s*\)$', cell_text):
                                            is_eq = True
                                            break
                                if is_eq:
                                    break
                    table_eq_cache[tbl_id] = is_eq

                if not is_eq:
                    in_table = True
                    tbl_id_ref = id(tbl)
                    p_row = next(p_el.iterancestors(w_tr), None)
                    if p_row is not None:
                        # ═══ DETECCIÓN ROBUSTA DEL HEADER ═══
                        # 1) Si la fila está en el conjunto de filas de header (incluye
                        #    múltiples filas marcadas con <w:tblHeader/>) → header
                        # 2) Si la celda padre del párrafo es continuación de un vMerge
                        #    cuyo restart está en una fila de header → header
                        # 3) Fallback: si es la primera fila XML → header
                        header_rows = self.header_rows_by_tbl.get(id(tbl), set())
                        vmerge_cells = self.vmerge_header_cells.get(id(tbl), set())

                        if id(p_row) in header_rows:
                            is_table_header = True
                        else:
                            # Verificar si la celda padre es vMerge continuation del header
                            p_cell = next(p_el.iterancestors(w_tc), None)
                            if p_cell is not None and id(p_cell) in vmerge_cells:
                                is_table_header = True
                            else:
                                # Fallback estricto: primera fila XML
                                first_row = next(tbl.iter(w_tr), None)
                                if p_row == first_row:
                                    is_table_header = True

            # Determinar si es un título (manual o automático de Word)
            is_heading = False
            heading_level = None

            hierarchy_match = re.match(r"^(\d+(?:\.\d+)+)\.?(?:[\s\t]+(.*))?$", txt.strip())
            if hierarchy_match and not in_table:
                is_heading = True
                heading_level = hierarchy_match.group(1).count('.') + 1

            num_pr = ppr_el.find('w:numPr', NSMAP) if ppr_el is not None else None
            is_heading_style = bool(re.search(r'(heading|ttulo|titulo)', style_id, re.IGNORECASE))

            if is_heading_style and not in_table:
                is_heading = True
                if num_pr is not None:
                    ilvl_el = num_pr.find('w:ilvl', NSMAP)
                    if ilvl_el is not None:
                        heading_level = int(ilvl_el.get(f'{{{NSMAP["w"]}}}val')) + 1

                if heading_level is None:
                    m_style = re.search(r'(heading|ttulo|titulo)\s*(\d+)', style_id, re.IGNORECASE)
                    if m_style:
                        style_num = int(m_style.group(2))
                        if style_num == 1:
                            dec_match = re.match(r'^(\d+(?:\.\d+)+)', txt.strip())
                            if dec_match and dec_match.group(1).count('.') >= 1:
                                heading_level = 2
                            else:
                                heading_level = 1
                        elif style_num == 2:
                            heading_level = 2
                        elif style_num == 3:
                            heading_level = 3
                        elif style_num == 4:
                            heading_level = 4
                        elif style_num == 5:
                            heading_level = 5
                        else:
                            heading_level = style_num

            if not is_heading and num_pr is not None and not in_table:
                ilvl_el = num_pr.find('w:ilvl', NSMAP)
                if ilvl_el is not None:
                    ilvl_val = int(ilvl_el.get(f'{{{NSMAP["w"]}}}val'))
                    inferred_level = ilvl_val + 1
                    if len(txt.strip()) < 200 and inferred_level >= 2:
                        is_heading = True
                        if heading_level is None:
                            heading_level = inferred_level

            if is_heading and heading_level is not None:
                self.current_level = heading_level

            # Resolviendo propiedades de lista
            num_id_val = None
            ilvl_val = None
            list_fmt = None
            list_lvl_text = None
            list_font = None
            if num_pr is not None:
                numId_el = num_pr.find('w:numId', NSMAP)
                ilvl_el = num_pr.find('w:ilvl', NSMAP)
                if numId_el is not None:
                    num_id_val = int(numId_el.get(f'{{{NSMAP["w"]}}}val'))
                if ilvl_el is not None:
                    ilvl_val = int(ilvl_el.get(f'{{{NSMAP["w"]}}}val'))
                
                if num_id_val is not None and ilvl_val is not None:
                    fmt_info = bullet_definitions.get((num_id_val, ilvl_val))
                    if fmt_info:
                        list_fmt = fmt_info['numFmt']
                        list_lvl_text = fmt_info['lvlText']
                        list_font = fmt_info['font']

            has_hyperlink = (p_el.find('.//w:hyperlink', NSMAP) is not None)

            self.paragraphs.append({
                "text": txt,
                "norm": norm_txt,
                "alignment": res_ppr.get('alignment', 'left'),
                "line_spacing": res_ppr.get('line_spacing'),
                "indent_first": res_ppr.get('indent_first'),
                "spacing_after": res_ppr.get('spacing_after', 0),
                "spacing_before": res_ppr.get('spacing_before', 0),
                "runs": runs,
                "index": idx,
                "estimated_page": estimated_page,
                "section": current_section,
                "has_page_break": has_page_break,
                "has_real_page_marker": has_real_page_marker,
                "style_id": style_id,
                "in_table": in_table,
                "is_table_header": is_table_header,
                "level": self.current_level,
                "is_heading": is_heading,
                "indent_left": res_ppr.get('indent_left'),
                "indent_hanging": res_ppr.get('indent_hanging'),
                "drawings": drawings,
                "has_hyperlink": has_hyperlink,
                "num_id": num_id_val,
                "ilvl": ilvl_val,
                "list_fmt": list_fmt,
                "list_lvl_text": list_lvl_text,
                "list_font": list_font,
                "has_omml": has_omml,
                "is_display_equation": is_display_equation,
                "tbl_id": tbl_id_ref,
                "has_textbox": has_textbox,
                "textbox_text": textbox_text,
                "looks_like_screenshot": looks_like_screenshot,
                "screenshot_drawing": screenshot_drawing,
            })

        # Identificar la portada (primera página)
        in_cover = True
        for idx, p in enumerate(self.paragraphs):
            txt = p["text"].strip()
            norm = p["norm"]

            if not in_cover:
                p["is_cover"] = False
                continue

            if norm in ["DEDICATORIA", "AGRADECIMIENTOS", "INDICE GENERAL", "RESUMEN", "ABSTRACT", "INTRODUCCION"]:
                in_cover = False
                p["is_cover"] = False
                continue

            p["is_cover"] = in_cover

            if p.get("has_page_break") or p.get("estimated_page", 1) > 1:
                in_cover = False

        # Encontrar el final del bloque de índices de forma totalmente robusta (max_index_idx)
        self.index_start_idx = -1
        max_index_idx = -1
        
        in_index_block = False
        gap_count = 0
        
        for idx, p_data in enumerate(self.paragraphs):
            norm_txt = p_data["norm"]
            txt = p_data["text"].strip()
            style = p_data.get("style_id", "")
            
            # Si encontramos el inicio real del cuerpo (no una línea de índice), terminamos la búsqueda
            is_index_line = "...." in txt or bool(re.search(r"\s+\d+$", txt)) or style.upper().startswith("TOC") or style.upper().startswith("TDC")
            if not is_index_line and ("CAPITULO" in norm_txt or "INTRODUCCION" in norm_txt):
                break

            if not in_index_block:
                is_index_title = norm_txt in [
                    "INDICE GENERAL", "INDICE DE TABLAS", "INDICE DE FIGURAS", 
                    "INDICE DE CUADROS", "INDICE DE ILUSTRACIONES", "INDICE DE ANEXOS", 
                    "INDICE DE GRAFICOS", "INDICE", "CONTENIDO", "TABLA DE MATERIAS",
                    "TABLA DE CONTENIDOS", "TABLA DE CONTENIDO"
                ] or (norm_txt.startswith("INDICE") and len(norm_txt) < 40)
                is_tdc = style.upper().startswith("TOC") or style.upper().startswith("TDC") if style else False
                
                if is_index_title or is_tdc:
                    if self.index_start_idx == -1:
                        self.index_start_idx = idx
                    in_index_block = True
                    gap_count = 0
                    max_index_idx = idx
                    continue
                
            if in_index_block:
                is_tdc = style.upper().startswith("TOC") or style.upper().startswith("TDC") if style else False
                is_idx_line = "...." in txt or bool(re.search(r"\s+\d+$", txt)) or is_tdc
                
                if is_idx_line or is_tdc:
                    max_index_idx = idx
                    gap_count = 0
                else:
                    gap_count += 1
                    
                # Si pasamos 10 párrafos sin características de índice, salimos del bloque de índice pero seguimos buscando
                if gap_count > 10:
                    in_index_block = False
                    gap_count = 0
                    
        self.last_index_idx = max_index_idx

        # Encontrar el inicio del cuerpo del documento (después del índice)
        body_start_idx = len(self.paragraphs)
        for idx, p_data in enumerate(self.paragraphs):
            if idx <= self.last_index_idx:
                continue
                
            norm_txt = p_data["norm"]
            is_capitulo = bool(re.match(r"^CAP[ÍI]TULO\s+(I|V|X|L|C|[0-9]+)", norm_txt))
            if is_capitulo or "INTRODUCCION" in norm_txt:
                body_start_idx = idx
                break

        # ── Pre-computar body_level y is_in_body para cada párrafo ──
        # Esto permite que los auditores por nivel filtren sin duplicar lógica de rastreo
        current_body_level = 1
        in_body = False
        for i, p in enumerate(self.paragraphs):
            txt = p["text"].strip()
            norm = p["norm"]
            sec_upper = p.get('section', '').upper()

            is_prelim_section = any(k in sec_upper for k in [
                'ÍNDICE', 'INDICE', 'DEDICATORIA', 'AGRADECIMIENTO', 'ACRÓNIMO', 'ACRONIMO', 'RESUMEN', 'ABSTRACT',
                'CONTENIDO', 'TABLA DE CONTENIDOS', 'TABLA DE CONTENIDO'
            ]) or any(k in norm for k in [
                'ÍNDICE GENERAL', 'INDICE GENERAL', 'ÍNDICE DE TABLAS', 'INDICE DE TABLAS', 'ÍNDICE DE FIGURAS', 'INDICE DE FIGURAS', 'ÍNDICE DE ANEXOS', 'INDICE DE ANEXOS',
                'CONTENIDO', 'TABLA DE CONTENIDOS', 'TABLA DE CONTENIDO'
            ])

            # Determinar si estamos en el cuerpo del documento
            if (self.last_index_idx != -1 and i <= self.last_index_idx) or is_prelim_section:
                p["is_in_body"] = False
                p["body_level"] = 0
                continue

            is_index_line = "...." in txt or bool(re.search(r"\d+$", txt))
            if not is_index_line and ("CAPITULO" in norm or "INTRODUCCION" in norm):
                in_body = True
                current_body_level = 1

            if not in_body or p.get('in_table') or (self.anexos_start_idx != -1 and i >= self.anexos_start_idx):
                p["is_in_body"] = False
                p["body_level"] = 0
                continue

            p["is_in_body"] = True

            # Detectar viñetas
            is_bullet = bool(txt) and (txt[0] in '-\u2022*•' or (len(txt) >= 1 and txt[0] == ''))
            p["is_bullet"] = is_bullet

            # Actualizar nivel basado en títulos
            is_capitulo = bool(re.match(r"^CAP[ÍI]TULO\s+(I|V|X|L|C|[0-9]+)", norm))
            es_seccion_principal = any(k in norm for k in [
                "INTRODUCCION", "MARCO TEORICO", "METODOLOGIA",
                "MATERIALES Y METODOS", "RESULTADOS Y DISCUSION",
                "CONCLUSIONES", "RECOMENDACIONES", "REFERENCIAS BIBLIOGRAFICAS"
            ]) and (p.get('style_id', '').upper().startswith('HEADING') or txt.isupper())

            numbering_match = re.match(r'^(\d+(?:\.\d+)+)\.?(?:[\s\t]+|\S)', txt)
            if is_capitulo or es_seccion_principal:
                current_body_level = 1
            elif numbering_match:
                current_body_level = numbering_match.group(1).count('.') + 1
            elif p.get('is_heading') and p.get('level'):
                current_body_level = p['level']

            p["body_level"] = current_body_level

        # Registrar secciones encontradas
        for p in self.paragraphs:
            norm = p["norm"]
            for kw in SECTION_KEYWORDS:
                if norm == kw:
                    self.sections_found.add(norm)
                    break
            if re.match(r"^CAPITULO\s+(I|V|X|L|C|[0-9]+)", norm):
                self.sections_found.add("CAPITULO")

    # ══════════════════════════════════════════════════════════════════════
    # UTILIDADES COMPARTIDAS
    # ══════════════════════════════════════════════════════════════════════

    def _norm(self, t):
        if not t: return ""
        return unicodedata.normalize('NFD', t).encode('ascii', 'ignore').decode('ascii').upper().strip()

    def _norm_alphanumeric(self, t):
        n = self._norm(t)
        return re.sub(r'[^A-Z0-9]', '', n)

    def _add(self, cat, rule, status, msg, expected="", actual="", p_idx=None, p_text="", page=None, section=None):
        if status == "error": self.stats["errors"] += 1
        elif status == "warning": self.stats["warnings"] += 1
        if p_idx is not None and p_idx < len(self.paragraphs):
            p_data = self.paragraphs[p_idx]
            if page is None:
                page = p_data.get("estimated_page")
            if section is None:
                section = p_data.get("section")
        self.results.append({
            "category": cat, "rule": rule, "status": status, "message": msg, 
            "expected": str(expected), "actual": str(actual),
            "paragraphIndex": p_idx, "paragraphText": p_text,
            "page": page,
            "section": section or "General",
        })

    def _get_p_props(self, p):
        if not p["runs"]: return 12, False, False, "Times New Roman"
        for r in p["runs"]:
            if r["text"].strip():
                return r["size"], r["bold"], r["italic"], r["font"]
        return 12, False, False, "Times New Roman"

    def extract_metadata(self):
        """
        Extrae automáticamente metadatos de la portada para alimentar el sistema de
        aprendizaje (ai_engine): facultad, escuela profesional, año, tipo de grado.

        Retorna un diccionario; los valores no detectados son None.
        """
        meta = {
            "faculty": None,
            "school": None,
            "year": None,
            "degree_type": None,
        }
        cover = [p for p in self.paragraphs if p.get("is_cover")]
        for p in cover:
            txt = p["text"].strip()
            if not txt:
                continue
            norm = p["norm"]

            # Facultad
            if meta["faculty"] is None:
                m = re.match(r"^FACULTAD\s+DE\s+(.+?)$", txt, re.IGNORECASE)
                if m:
                    meta["faculty"] = m.group(1).strip().rstrip(",.").title()

            # Escuela profesional
            if meta["school"] is None:
                m = re.match(r"^ESCUELA\s+PROFESIONAL\s+DE\s+(.+?)$", txt, re.IGNORECASE)
                if m:
                    meta["school"] = m.group(1).strip().rstrip(",.").title()

            # Año (4 dígitos sueltos)
            if meta["year"] is None and re.fullmatch(r"\d{4}", txt):
                year_val = int(txt)
                if 2000 <= year_val <= 2100:
                    meta["year"] = year_val

            # Tipo de grado
            if meta["degree_type"] is None:
                if "LICENCIADO" in norm or "INGENIERO" in norm or "BACHILLER" in norm:
                    meta["degree_type"] = "Pregrado"
                elif "MAESTRO" in norm or "MAESTR" in norm:
                    meta["degree_type"] = "Maestría"
                elif "DOCTOR" in norm:
                    meta["degree_type"] = "Doctorado"

        return meta

    def _finalize(self):
        total = len(self.results)
        passed = sum(1 for r in self.results if r["status"] == "passed")
        self.stats.update({"passed": passed, "score": max(0, round((passed/total)*100)) if total > 0 else 0})
        return {"stats": self.stats, "results": self.results}
